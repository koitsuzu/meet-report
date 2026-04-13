#!/usr/bin/env python3
"""
generate_docs.py — meeting_summary skill

Takes a JSON file conforming to references/schema.json and produces:
  - <output_stem>.docx  (Word document matching the MIS meeting template)
  - <output_stem>.md    (Markdown version)

Usage:
    python generate_docs.py --input meeting.json --output 會議記錄_20260326
    python generate_docs.py --input meeting.json   # uses datetime from JSON as stem
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)


# ─────────────────────────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True, color="AAAAAA"):
    """Apply single-line border to specific sides of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, enabled in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        if enabled:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_paragraph(cell, text: str, bold=False, font_size=10, color=None, indent=False):
    """Add a clean styled paragraph to a table cell."""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# DOCX generation
# ─────────────────────────────────────────────────────────────────────────────

HEADER_BG = "1F497D"   # dark blue (MIS template header colour)
HEADER_FG = "FFFFFF"
SUBHEADER_BG = "D9E1F2"  # light blue
ROW_ALT_BG = "F2F2F2"


def build_docx(data: dict, output_path: Path):
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Document title ────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("會議記錄")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(HEADER_BG)

    doc.add_paragraph()  # spacer

    # ── Meeting info table (2 rows × 4 cols) ─────────────────────────────────
    info = data["meeting_info"]
    attendees = data["attendees"]

    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = "Table Grid"

    # Row 0: 會議主題 | value | 會議室 | value
    def _info_row(row_idx, label1, val1, label2, val2):
        row = info_table.rows[row_idx]
        cells = row.cells
        # Merge pairs
        for i, (label, val) in enumerate([(label1, val1), (label2, val2)]):
            lc = cells[i * 2]
            vc = cells[i * 2 + 1]
            set_cell_bg(lc, HEADER_BG)
            lc.paragraphs[0].clear()
            lr = lc.paragraphs[0].add_run(label)
            lr.bold = True
            lr.font.color.rgb = RGBColor.from_string(HEADER_FG)
            lr.font.size = Pt(10)
            vc.paragraphs[0].clear()
            vr = vc.paragraphs[0].add_run(str(val))
            vr.font.size = Pt(10)

    _info_row(0, "會議主題", info["topic"],    "會議室", info["room"])
    _info_row(1, "日期時間：", info["datetime"], "會議記錄", info["recorder"])

    # Row 2: 應出席人員 (merged across 3 cols)
    r2 = info_table.rows[2]
    set_cell_bg(r2.cells[0], HEADER_BG)
    r2.cells[0].paragraphs[0].clear()
    run = r2.cells[0].paragraphs[0].add_run("應出席人員")
    run.bold = True; run.font.color.rgb = RGBColor.from_string(HEADER_FG); run.font.size = Pt(10)
    r2.cells[1].merge(r2.cells[2]).merge(r2.cells[3])
    r2.cells[1].paragraphs[0].clear()
    r2.cells[1].paragraphs[0].add_run(", ".join(attendees["expected"])).font.size = Pt(10)

    # Row 3: 與會人員 (merged across 3 cols)
    r3 = info_table.rows[3]
    set_cell_bg(r3.cells[0], HEADER_BG)
    r3.cells[0].paragraphs[0].clear()
    run = r3.cells[0].paragraphs[0].add_run("與會人員")
    run.bold = True; run.font.color.rgb = RGBColor.from_string(HEADER_FG); run.font.size = Pt(10)
    r3.cells[1].merge(r3.cells[2]).merge(r3.cells[3])
    r3.cells[1].paragraphs[0].clear()
    r3.cells[1].paragraphs[0].add_run(", ".join(attendees["present"])).font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Discussion content table ──────────────────────────────────────────────
    disc_table = doc.add_table(rows=1, cols=1)
    disc_table.style = "Table Grid"

    # Section header
    hdr_row = disc_table.rows[0]
    set_cell_bg(hdr_row.cells[0], HEADER_BG)
    hdr_row.cells[0].paragraphs[0].clear()
    h = hdr_row.cells[0].paragraphs[0].add_run("討論內容")
    h.bold = True; h.font.color.rgb = RGBColor.from_string(HEADER_FG); h.font.size = Pt(11)

    # Sub-header: week label
    sub_row = disc_table.add_row()
    set_cell_bg(sub_row.cells[0], SUBHEADER_BG)
    sub_row.cells[0].paragraphs[0].clear()
    sub_row.cells[0].paragraphs[0].add_run("本週：").font.size = Pt(10)

    # Each speaker
    for entry in data.get("discussion", []):
        row = disc_table.add_row()
        cell = row.cells[0]
        # Clear default paragraph
        cell.paragraphs[0].clear()
        spk_run = cell.paragraphs[0].add_run(f"{entry['speaker']}：")
        spk_run.bold = True
        spk_run.font.size = Pt(10)
        for point in entry.get("points", []):
            bp = cell.add_paragraph()
            bp.paragraph_format.left_indent = Cm(0.8)
            br = bp.add_run(f"• {point}")
            br.font.size = Pt(10)

    # Announcements (if any)
    announcements = data.get("announcements", [])
    if announcements:
        ann_row = disc_table.add_row()
        cell = ann_row.cells[0]
        cell.paragraphs[0].clear()
        for ann in announcements:
            p = cell.paragraphs[0] if ann == announcements[0] else cell.add_paragraph()
            p.add_run(ann).font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Action items table ──────────────────────────────────────────────────
    action_items = data.get("action_items", [])
    COLS = ["項目", "說明", "負責人", "狀態", "預定完成日"]
    COL_WIDTHS = [Cm(2), Cm(6), Cm(2.5), Cm(2.5), Cm(2.5)]

    act_table = doc.add_table(rows=1, cols=5)
    act_table.style = "Table Grid"

    # Header row
    hdr = act_table.rows[0]
    for i, (col, w) in enumerate(zip(COLS, COL_WIDTHS)):
        cell = hdr.cells[i]
        set_cell_bg(cell, HEADER_BG)
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(col)
        r.bold = True; r.font.color.rgb = RGBColor.from_string(HEADER_FG); r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if action_items:
        for idx, ai in enumerate(action_items):
            row = act_table.add_row()
            if idx % 2 == 1:
                for c in row.cells:
                    set_cell_bg(c, ROW_ALT_BG)
            vals = [ai.get("item",""), ai.get("description",""), ai.get("owner",""),
                    ai.get("status",""), ai.get("due_date","")]
            for i, (cell, val, w) in enumerate(zip(row.cells, vals, COL_WIDTHS)):
                cell.width = w
                cell.paragraphs[0].clear()
                cell.paragraphs[0].add_run(str(val)).font.size = Pt(10)
    else:
        # Add a blank placeholder row
        row = act_table.add_row()
        row.cells[0].paragraphs[0].add_run("1.").font.size = Pt(10)

    # ── Footer note ────────────────────────────────────────────────────────────
    note = data.get("facilitator_note", "")
    if note:
        doc.add_paragraph()
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.add_run(note).font.size = Pt(9)

    doc.save(str(output_path))
    print(f"✅  Word saved → {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# Markdown generation
# ─────────────────────────────────────────────────────────────────────────────

def build_markdown(data: dict, output_path: Path):
    info = data["meeting_info"]
    attendees = data["attendees"]
    lines = []

    lines.append(f"# 會議記錄")
    lines.append("")
    lines.append(f"| 項目 | 內容 |")
    lines.append(f"|------|------|")
    lines.append(f"| 會議主題 | {info['topic']} |")
    lines.append(f"| 會議室   | {info['room']} |")
    lines.append(f"| 日期時間 | {info['datetime']} |")
    lines.append(f"| 會議記錄 | {info['recorder']} |")
    lines.append(f"| 應出席   | {', '.join(attendees['expected'])} |")
    lines.append(f"| 與會人員 | {', '.join(attendees['present'])} |")
    lines.append("")
    lines.append("## 討論內容")
    lines.append("")
    lines.append("**本週：**")
    lines.append("")
    for entry in data.get("discussion", []):
        lines.append(f"**{entry['speaker']}：**")
        for point in entry.get("points", []):
            lines.append(f"  - {point}")
        if not entry.get("points"):
            lines.append("  - （無）")
        lines.append("")

    announcements = data.get("announcements", [])
    if announcements:
        for ann in announcements:
            lines.append(f"> {ann}")
        lines.append("")

    lines.append("## 追蹤項目")
    lines.append("")
    lines.append("| 項目 | 說明 | 負責人 | 狀態 | 預定完成日 |")
    lines.append("|------|------|--------|------|------------|")
    for ai in data.get("action_items", []):
        lines.append(f"| {ai.get('item','')} | {ai.get('description','')} | {ai.get('owner','')} | {ai.get('status','')} | {ai.get('due_date','')} |")

    note = data.get("facilitator_note", "")
    if note:
        lines.append("")
        lines.append(f"---")
        lines.append(f"*{note}*")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"✅  Markdown saved → {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate meeting minutes from structured JSON.")
    parser.add_argument("--input",  "-i", required=True, help="Path to JSON file (schema.json format)")
    parser.add_argument("--output", "-o", default=None,  help="Output stem (without extension). Defaults to datetime from JSON.")
    parser.add_argument("--dir",    "-d", default=".",   help="Output directory (default: current dir)")
    args = parser.parse_args()

    json_path = Path(args.input)
    if not json_path.exists():
        print(f"ERROR: Input file not found: {json_path}")
        sys.exit(1)

    data = json.loads(json_path.read_text(encoding="utf-8"))

    # Determine output stem
    if args.output:
        stem = args.output
    else:
        dt_str = data.get("meeting_info", {}).get("datetime", "")
        try:
            dt = datetime.strptime(dt_str, "%Y/%m/%d %H:%M")
            stem = f"會議記錄_{dt.strftime('%Y%m%d')}_{data['meeting_info'].get('topic','Meeting')}"
        except Exception:
            stem = "會議記錄_output"

    out_dir = Path(args.dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    build_docx(data,     out_dir / f"{stem}.docx")
    build_markdown(data, out_dir / f"{stem}.md")


if __name__ == "__main__":
    main()
